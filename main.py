import os
import random
import streamlit as st
import pandas as pd
from datetime import datetime
from groq import Groq
from docx import Document
from PyPDF2 import PdfReader

# Initialize the Groq client
api_key = "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"
if not api_key:
    raise ValueError("API key is missing")
client = Groq(api_key=api_key)

def groq_response(prompt):
    try:
        chat_completion = client.chat.completions.create(
            messages=[
                {"role": "system", "content": "You are an intelligent project proposal assistant chatbot, designed to help create detailed and structured project proposals by extracting key information and guiding users through technical aspects and planning."},
                {"role": "user", "content": prompt}
            ],
            model="llama3-8b-8192"
        )
        return chat_completion.choices[0].message.content.strip()
    except Exception as e:
        return f"Error: {str(e)}"
# Add this new function after the groq_response function
def generate_project_overview(project_idea):
    prompt = f"""
    Based on this project idea:
    {project_idea}

    Generate a comprehensive project overview that includes:
    1. A clear explanation of what the project is
    2. The main problems it solves
    3. The target audience/beneficiaries
    4. The key value propositions
    5. The expected impact/benefits

    Format this as 2-3 well-structured paragraphs that flow naturally.
    Do not simply list points, but create a cohesive narrative about the project.
    """
    return groq_response(prompt)

# New function to extract core features/modules from project idea
def extract_core_features(project_idea):
    prompt = f"""
    Based on this project idea:
    {project_idea}
    
    Extract and list all core features/modules that would be essential for this project.
    Format as a comma-separated list of core modules/features.
    Focus on main functionalities, not technical implementation details.
    """
    features = groq_response(prompt)
    return features.split(',') if ',' in features else [features]

# New function to extract user roles and use cases from project idea
def extract_user_roles(project_idea):
    prompt = f"""
    Based on this project idea:
    {project_idea}
    
    Identify all user roles that would interact with this system.
    Format as a comma-separated list of roles.
    Include all stakeholders who would use the system.
    """
    roles = groq_response(prompt)
    return roles.split(',') if ',' in roles else [roles]

def generate_wbs_and_estimation(project_name, modules, tech_stack, complexity, overlap, roadmap_basis):
    prompt = f"""
    Project Proposal - {project_name}
    - Features: {modules}
    - Tech Stack: {tech_stack}
    - Complexity: {complexity}
    - Overlap: {overlap}
    - Roadmap Basis: {roadmap_basis}
    Based on this, provide a detailed Work Breakdown Structure (WBS) including estimations for:
    - UI/UX Design
    - Frontend Development
    - Backend Development
    - Machine Learning (if applicable)
    - API Integration
    - QA and Testing
    - Deployment
    """
    return groq_response(prompt)

def generate_diagrams(user_roles):
    dfd_prompt = f"Generate a data flow diagram (DFD) for a system with these user roles: {user_roles}. Include data flows, processes, data stores, and external entities."
    erd_prompt = f"Generate an entity relationship diagram (ERD) for a system with these user roles: {user_roles}. Focus on database entities, attributes, and relationships."
    
    user_flow = groq_response(f"Generate a user flow in Mermaid.js for the following roles: {user_roles}")
    dfd = groq_response(dfd_prompt)
    erd = groq_response(erd_prompt)
    return user_flow, dfd, erd

def generate_user_stories_and_use_cases(modules, user_roles):
    prompt = f"""
    Based on the following project modules and user roles:
    - Modules: {modules}
    - User Roles: {user_roles}
    
    Generate user stories and use cases in a structured format with these columns:
    Role | User Story | Use Case
    
    Ensure each story follows "As a [Role], I can [Action] so that [Goal]" format.
    """
    response = groq_response(prompt)
    
    stories = []
    for line in response.split('\n'):
        if 'As a' in line:
            try:
                role = line.split('As a')[1].split(',')[0].strip()
                action = line.split('I can')[1].split('so that')[0].strip()
                goal = line.split('so that')[1].strip()
                module = next((m for m in modules if m.lower() in action.lower()), "General")
                user_story = f"As a {role}, I can {action} so that {goal}"
                use_case = f"{role} can perform the action: {action} to achieve: {goal}"
                stories.append({
                    "Role": role,
                    "User Story": user_story,
                    "Use Case": use_case
                })
            except:
                continue
    return stories

def calculate_timeline(complexity):
    if complexity == "Low":
        return random.randint(12, 18)
    elif complexity == "Medium":
        return random.randint(18, 36)
    elif complexity == "High":
        return random.randint(30, 52)

def extract_text_from_file(uploaded_file):
    if uploaded_file.type == "application/pdf":
        pdf_reader = PdfReader(uploaded_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    return None

def generate_cost_estimation(wbs, complexity):
    base_costs = {
        "UI/UX": {"Design": 1000, "Prototyping": 500, "User Testing": 300},
        "Frontend": {"Development": 2000, "Integration": 500, "Optimization": 300},
        "Backend": {"API Development": 2000, "Database": 1000, "Security": 800},
        "ML": {"Model Development": 3000, "Training": 1500, "Integration": 1000},
        "API Integration": {"External APIs": 1000, "Documentation": 500},
        "QA": {"Testing": 800, "Automation": 400},
        "Deployment": {"Setup": 500, "Configuration": 300}
    }
    
    complexity_multiplier = {
        "Low": 1.0,
        "Medium": 1.25,
        "High": 1.5
    }
    
    detailed_costs = []
    for category, tasks in base_costs.items():
        category_total = 0
        for task, cost in tasks.items():
            adjusted_cost = cost * complexity_multiplier[complexity]
            category_total += adjusted_cost
            detailed_costs.append({
                "Category": category,
                "Task": task,
                "Base Cost": cost,
                "Adjusted Cost": round(adjusted_cost, 2)
            })
    
    return detailed_costs

def save_to_word(project_name, project_overview, modules, tech_stack, user_roles, user_stories, wbs, user_flow, dfd, erd, cost_estimation):
    doc = Document()
    doc.add_heading(f'Project Proposal: {project_name}', 0)

    doc.add_heading('Project Overview', level=1)
    doc.add_paragraph(project_overview)

    doc.add_heading('Modules / Features', level=1)
    doc.add_paragraph(", ".join(modules))

    doc.add_heading('Tech Stacks', level=1)
    doc.add_paragraph(tech_stack)

    doc.add_heading('User Roles', level=1)
    doc.add_paragraph(", ".join(user_roles))

    doc.add_heading('User Stories and Use Cases', level=1)
    doc.add_paragraph(str(user_stories))

    doc.add_heading('User Flow (Mermaid.js)', level=1)
    doc.add_paragraph(user_flow)

    doc.add_heading('DFD / Backend Architecture (Mermaid.js)', level=1)
    doc.add_paragraph(dfd)

    doc.add_heading('ERD (Entity Relationship Diagram) (Mermaid.js)', level=1)
    doc.add_paragraph(erd)

    doc.add_heading('Work Breakdown Structure (WBS)', level=1)
    doc.add_paragraph(wbs)

    doc.add_heading('Cost Estimation', level=1)
    doc.add_paragraph(str(cost_estimation))

    doc.add_heading('Disclaimer', level=1)
    doc.add_paragraph("This project proposal is valid for 30 days from the date of issue. Please review and confirm all details within this period.")

    file_path = f"{project_name}_Proposal_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    doc.save(file_path)
    return file_path

def display_project_proposal(project_name, project_overview, modules, tech_stack, user_roles, user_stories, wbs, user_flow, dfd, erd, cost_estimation):
    st.markdown(f"### Title: {project_name}")
    st.markdown(f"### Date: {datetime.now().strftime('%Y-%m-%d')}")
    st.markdown("### Reviewer: Mr. Aqib")

    st.markdown("### Project Overview")
    st.write(project_overview)

    st.markdown("### Extracted Modules / Features")
    st.write(", ".join(modules))

    st.markdown("### Tech Stacks")
    st.write(tech_stack)

    st.markdown("### Extracted User Roles")
    st.write(", ".join(user_roles))

    st.markdown("### User Stories and Use Cases")
    if user_stories:
        df_stories = pd.DataFrame(user_stories)
        st.table(df_stories)

    st.markdown("### User Flow (Mermaid.js)")
    st.code(user_flow, language="mermaid")

    st.markdown("### DFD / Backend Architecture (Mermaid.js)")
    st.code(dfd, language="mermaid")

    st.markdown("### ERD (Entity Relationship Diagram) (Mermaid.js)")
    st.code(erd, language="mermaid")

    st.markdown("### Work Breakdown Structure (WBS)")
    st.write(wbs)

    st.markdown("### Cost Estimation")
    df_costs = pd.DataFrame(cost_estimation)
    total_cost = df_costs['Adjusted Cost'].sum()
    
    st.table(df_costs)
    st.markdown(f"**Total Project Cost Estimate: ${total_cost:,.2f}**")

    word_file_path = save_to_word(project_name, project_overview, modules, tech_stack, user_roles, 
                                 pd.DataFrame(user_stories).to_string(), wbs, user_flow, dfd, erd, 
                                 pd.DataFrame(cost_estimation).to_string())
    with open(word_file_path, "rb") as f:
        st.download_button(
            label="Download Proposal as Word Document",
            data=f,
            file_name=word_file_path,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


def main():
    st.title("Project Proposal Chatbot")

    with st.form(key="project_form"):
        project_name = st.text_input("Enter the project title", value="My Project")
        project_idea = st.text_area("Provide the overall project idea")  # Renamed from project_overview to project_idea
        tech_stack = st.text_input("Do you have a preferred tech stack? (Leave blank for suggestions)")

        design = st.radio("Do you have the design?", ("Yes", "No"))
        overlap = None
        if design == "Yes":
            overlap = st.radio("Do you want FE/BE overlap?", ("Yes", "No"))

        roadmap = st.radio("Do you need the Roadmap?", ("Yes", "No"))
        roadmap_basis = None
        if roadmap == "Yes":
            roadmap_basis = st.radio("Specify the roadmap basis", ("Week-based", "Month-based"))
        else:
            roadmap_basis = "No roadmap"

        complexity = st.selectbox("Select project complexity", ["Low", "Medium", "High"])

        uploaded_file = st.file_uploader("Upload project requirement file (PDF, DOC)", type=["pdf", "docx"])

        if uploaded_file:
            extracted_text = extract_text_from_file(uploaded_file)
            st.text_area("Extracted Text", extracted_text, height=300)

        if not project_name or not project_idea:
            st.error("Project name and project idea are required fields!")

        submitted = st.form_submit_button("Generate Proposal")

        if submitted:
            st.info("Generating proposal... This may take a moment.")

            # Generate proper project overview from the idea
            project_overview = generate_project_overview(project_idea)

            # Extract modules and user roles from project idea
            modules = extract_core_features(project_idea)
            user_roles = extract_user_roles(project_idea)

            # Generate user stories and use cases using extracted information
            user_stories = generate_user_stories_and_use_cases(modules, user_roles)

            # Calculate timeline based on complexity
            total_timeline = calculate_timeline(complexity)

            # Generate WBS and diagrams
            wbs = generate_wbs_and_estimation(
                project_name,
                ", ".join(modules),
                tech_stack or "Tech stack not specified; recommend.",
                complexity,
                overlap if overlap else "No FE/BE overlap",
                roadmap_basis
            )
            user_flow, dfd, erd = generate_diagrams(user_roles)

            # Generate cost estimation
            cost_estimation = generate_cost_estimation(wbs, complexity)

            # Display the project proposal
            display_project_proposal(
                project_name,
                project_overview,  # Now using the generated overview instead of the raw idea
                modules,
                tech_stack,
                user_roles,
                user_stories,
                wbs,
                user_flow,
                dfd,
                erd,
                cost_estimation
            )
if __name__ == "__main__":
    main()